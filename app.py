from flask import Flask, render_template, request, jsonify, send_file, make_response
import os
import sys
import tempfile
from werkzeug.utils import secure_filename
import extract_employee_shifts
import shutil
from docx import Document
import re
import time
from io import BytesIO
try:
    from PIL import Image, ImageDraw, ImageFont
except Exception:
    Image = None

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

def _resolve_template_dir():
    # When frozen by py2app, the executable is .../Contents/MacOS/<app>
    # Resources (with templates) live at ../Resources/templates
    try:
        if getattr(sys, 'frozen', False):
            mac_resources = os.path.abspath(os.path.join(os.path.dirname(sys.executable), '..', 'Resources', 'templates'))
            if os.path.isdir(mac_resources):
                return mac_resources
        # Fallback to local directory
        local_templates = os.path.join(BASE_DIR, 'templates')
        return local_templates
    except Exception:
        return os.path.join(BASE_DIR, 'templates')

TEMPLATE_DIR = _resolve_template_dir()
app = Flask(__name__, template_folder=TEMPLATE_DIR)
app.secret_key = 'your-secret-key-here'  # Change this in production

# Configuration
# For a frozen/packaged app (py2app / PyInstaller) use a user-writable directory
if getattr(sys, 'frozen', False):  # Running inside a bundled app
    DEFAULT_UPLOAD_ROOT = os.path.join(os.path.expanduser('~'), 'ShiftsAnalyzerUploads')
else:
    DEFAULT_UPLOAD_ROOT = 'uploads'

UPLOAD_FOLDER = DEFAULT_UPLOAD_ROOT
ALLOWED_EXTENSIONS = {'docx'}
MAX_CONTENT_LENGTH = 100 * 1024 * 1024  # 100MB max file size

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['MAX_CONTENT_LENGTH'] = MAX_CONTENT_LENGTH

# Ensure upload folder exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def get_user_downloads_dir():
    """Return the user's Downloads directory (macOS/Linux/Windows simple heuristic)."""
    home = os.path.expanduser('~')
    # macOS & most Linux distros
    downloads = os.path.join(home, 'Downloads')
    if os.path.isdir(downloads):
        return downloads
    # Windows fallback
    windows_downloads = os.path.join(home, 'Downloads')
    return windows_downloads if os.path.isdir(windows_downloads) else home

@app.route('/test')
def test_page():
    return render_template('test.html')

@app.route('/test_route', methods=['GET', 'POST'])
def test_route():
    if request.method == 'POST':
        return jsonify({'status': 'POST received', 'data': dict(request.form)})
    return jsonify({'status': 'GET received'})

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/analyze', methods=['POST'])
def analyze_files():
    """Analyze uploaded files and return summary data for heatmap"""
    try:
        if 'files' not in request.files:
            return jsonify({'error': 'No files selected'}), 400
        
        files = request.files.getlist('files')
        
        if len(files) == 0:
            return jsonify({'error': 'No files selected'}), 400
        
        # Create a temporary directory for this analysis
        session_dir = tempfile.mkdtemp(dir=UPLOAD_FOLDER)
        
        # Save uploaded files with filename mapping
        uploaded_files = []
        filename_mapping = {}
        for file in files:
            if file and file.filename and allowed_file(file.filename):
                original_filename = file.filename
                secure_name = secure_filename(file.filename)
                file_path = os.path.join(session_dir, secure_name)
                file.save(file_path)
                
                filename_mapping[secure_name] = original_filename
                uploaded_files.append(secure_name)
        
        if not uploaded_files:
            shutil.rmtree(session_dir)
            return jsonify({'error': 'No valid .docx files uploaded'}), 400
        
        # Save filename mapping for later use
        mapping_file = os.path.join(session_dir, 'filename_mapping.txt')
        with open(mapping_file, 'w', encoding='utf-8') as f:
            for secure_name, original_name in filename_mapping.items():
                f.write(f'{secure_name}:{original_name}\n')
        
        # Analyze all employees and shifts
        summary_data = analyze_all_employees(session_dir, filename_mapping)
        
        return jsonify({
            'success': True,
            'session_dir': os.path.basename(session_dir),
            'summary': summary_data
        })
        
    except Exception as e:
        return jsonify({'error': f'Analysis failed: {str(e)}'}), 500

@app.route('/upload', methods=['POST'])
def upload_file():
    """Generate individual employee report"""
    try:
        employee_name = request.form.get('employee_name')
        session_id = request.form.get('session_id')
        
        if not employee_name or not session_id:
            return jsonify({'error': 'Missing employee name or session ID'}), 400
        
        session_dir = os.path.join(UPLOAD_FOLDER, session_id)
        
        if not os.path.exists(session_dir):
            return jsonify({'error': 'Session not found'}), 404
        
        # Load filename mapping
        mapping_file = os.path.join(session_dir, 'filename_mapping.txt')
        filename_mapping = {}
        
        if os.path.exists(mapping_file):
            with open(mapping_file, 'r', encoding='utf-8') as f:
                for line in f:
                    if ':' in line:
                        temp_name, original_name = line.strip().split(':', 1)
                        filename_mapping[temp_name] = original_name
        else:
            # If no mapping file, create one from current files
            for secure_fname in os.listdir(session_dir):
                if secure_fname.endswith('.docx'):
                    filename_mapping[secure_fname] = secure_fname
        
        # Extract shifts for the specific employee
        shifts = extract_with_mapping(employee_name, session_dir, filename_mapping)
        
        if not shifts:
            return jsonify({'error': f'No shifts found for employee: {employee_name}'}), 404
        
        # Generate Excel file
        output_file = os.path.join(session_dir, f'{employee_name}_shifts.xlsx')
        extract_employee_shifts.write_to_xlsx(shifts, output_file)
        # Ensure file fully written (occasionally needed in bundled mode)
        for _ in range(20):
            if os.path.exists(output_file) and os.path.getsize(output_file) > 0:
                break
            time.sleep(0.05)
        
        # Also copy the file to user's Downloads directory for immediate availability.
        auto_saved = False
        saved_path = ''
        try:
            downloads_dir = get_user_downloads_dir()
            os.makedirs(downloads_dir, exist_ok=True)
            base_target_name = f'{employee_name}_shifts.xlsx'
            target_path = os.path.join(downloads_dir, base_target_name)
            # Avoid overwriting existing file: add (1), (2), etc.
            if os.path.exists(target_path):
                stem, ext = os.path.splitext(base_target_name)
                counter = 1
                while True:
                    candidate = f"{stem} ({counter}){ext}"
                    candidate_path = os.path.join(downloads_dir, candidate)
                    if not os.path.exists(candidate_path):
                        target_path = candidate_path
                        break
                    counter += 1
            shutil.copyfile(output_file, target_path)
            auto_saved = True
            saved_path = target_path
        except Exception as copy_err:
            print(f"Could not auto-save to Downloads: {copy_err}")

        return jsonify({
            'success': True,
            'message': f'Found {len(shifts)} shifts for {employee_name}',
            'download_url': f'/download/{session_id}/{employee_name}_shifts.xlsx',
            'session_dir': session_id,
            'auto_saved': auto_saved,
            'saved_path': saved_path
        })
        
    except Exception as e:
        return jsonify({'error': f'Upload failed: {str(e)}'}), 500

def analyze_all_employees(session_dir, filename_mapping):
    """Analyze all employees and return summary data for heatmap"""
    employee_shifts = {}
    all_employee_names = set()
    
    # First pass: Extract all possible employee names to build a comprehensive list
    for secure_fname in os.listdir(session_dir):
        if not secure_fname.endswith('.docx'):
            continue
            
        filepath = os.path.join(session_dir, secure_fname)
        original_filename = filename_mapping.get(secure_fname, secure_fname)
        
        try:
            doc = Document(filepath)
            
            # Find the table and extract all possible names
            for table in doc.tables:
                for row in table.rows[1:]:  # Skip header
                    shift_type = row.cells[0].text.strip()
                    
                    # Skip "Assenti" entries
                    if 'assenti' in shift_type.lower():
                        continue
                    
                    for cell in row.cells[1:]:
                        cell_text = cell.text.strip()
                        if cell_text and not cell_text.isspace():
                            # Extract all possible names (both comma-separated and from parentheses)
                            # First, get names separated by commas/newlines
                            base_names = [name.strip() for name in cell_text.replace('\n', ',').split(',') if name.strip()]
                            
                            for base_name in base_names:
                                # Remove non-name parentheses content (like shifts, numbers, etc.)
                                # but preserve potential employee names
                                clean_base = re.sub(r'\([^)]*(?:turno|shift|ore|h|:|\d+)[^)]*\)', '', base_name, flags=re.IGNORECASE)
                                
                                # Extract main name (before parentheses)
                                main_name = re.sub(r'\([^)]*\)', '', clean_base).strip()
                                if main_name and len(main_name) > 1 and not is_roman_numeral(main_name):
                                    all_employee_names.add(main_name)
                                
                                # Extract potential names from parentheses
                                parentheses_matches = re.findall(r'\(([^)]+)\)', clean_base)
                                for match in parentheses_matches:
                                    potential_name = match.strip()
                                    # Check if it looks like a name (not a number, time, or shift info)
                                    if (potential_name and 
                                        len(potential_name) > 1 and 
                                        not re.match(r'^[\d:.-]+$', potential_name) and
                                        not re.search(r'\b(?:turno|shift|ore|h)\b', potential_name, re.IGNORECASE)):
                                        all_employee_names.add(potential_name)
        
        except Exception as e:
            print(f"Error in first pass processing {original_filename}: {e}")
            continue
    
    print(f"Found {len(all_employee_names)} unique employee names: {sorted(all_employee_names)}")
    
    # Second pass: Extract shifts using the comprehensive name list
    for secure_fname in os.listdir(session_dir):
        if not secure_fname.endswith('.docx'):
            continue
            
        filepath = os.path.join(session_dir, secure_fname)
        original_filename = filename_mapping.get(secure_fname, secure_fname)
        
        try:
            doc = Document(filepath)
            
            # Extract date range from ORIGINAL filename
            start_day, start_month, end_day, end_month, start_year, end_year = extract_employee_shifts.extract_date_range_from_filename(original_filename)
            if not all([start_day, start_month, end_day, end_month]):
                continue
            
            # Find the table and extract all employees
            for table in doc.tables:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                days = headers[1:]
                
                for row in table.rows[1:]:
                    shift_type = row.cells[0].text.strip()
                    
                    # Skip "Assenti" entries
                    if 'assenti' in shift_type.lower():
                        continue
                    
                    for i, cell in enumerate(row.cells[1:]):
                        cell_text = cell.text.strip()
                        if cell_text and not cell_text.isspace():
                            # Extract employee names using the comprehensive approach
                            employee_names = extract_employee_names_from_cell(cell_text, all_employee_names)
                            
                            for employee_name in employee_names:
                                if employee_name:
                                    if employee_name not in employee_shifts:
                                        employee_shifts[employee_name] = {}
                                    
                                    if shift_type not in employee_shifts[employee_name]:
                                        employee_shifts[employee_name][shift_type] = 0
                                    
                                    employee_shifts[employee_name][shift_type] += 1
                                    
                                    # Add weekend shifts for "Guardia" on Friday
                                    day_name = days[i] if i < len(days) else ''
                                    if 'guardia' in shift_type.lower() and day_name.lower() == 'venerdì':
                                        employee_shifts[employee_name][shift_type] += 2  # Saturday + Sunday
        
        except Exception as e:
            print(f"Error processing {original_filename}: {e}")
            continue
    
    return employee_shifts

def is_roman_numeral(text):
    """Check if text is a Roman numeral in parentheses that should be ignored"""
    if not text:
        return False
    
    # Remove parentheses if present
    clean_text = text.strip('()')
    
    # Check if it's a Roman numeral (I, II, III, IV, V, VI, VII, VIII, IX, X, etc.)
    roman_pattern = r'^(I{1,3}|IV|V|VI{0,3}|IX|X{1,3}|XL|L|LX{0,3}|XC|C{1,3}|CD|D|DC{0,3}|CM|M{1,3})$'
    return bool(re.match(roman_pattern, clean_text.upper()))

def extract_employee_names_from_cell(cell_text, known_names):
    """Extract employee names from a cell, considering both comma-separated and parentheses formats"""
    employee_names = []
    
    # Split by commas and newlines first
    base_parts = [part.strip() for part in cell_text.replace('\n', ',').split(',') if part.strip()]
    
    for part in base_parts:
        # Remove non-name parentheses content (shifts, times, etc.)
        cleaned_part = re.sub(r'\([^)]*(?:turno|shift|ore|h|:|\d+)[^)]*\)', '', part, flags=re.IGNORECASE)
        
        # Extract main name (before any parentheses)
        main_name = re.sub(r'\([^)]*\)', '', cleaned_part).strip()
        if main_name and main_name in known_names:
            employee_names.append(main_name)
        
        # Check parentheses for additional employee names
        parentheses_matches = re.findall(r'\(([^)]+)\)', cleaned_part)
        for match in parentheses_matches:
            potential_name = match.strip()
            
            # Skip Roman numerals
            if is_roman_numeral(potential_name):
                continue
                
            if potential_name in known_names:
                employee_names.append(potential_name)
            else:
                # Check if it's a partial match (like "Di Bella" matching "Di Bella")
                for known_name in known_names:
                    if (potential_name.lower() in known_name.lower() or 
                        known_name.lower() in potential_name.lower()) and len(potential_name) > 2:
                        employee_names.append(known_name)
                        break
    
    return list(set(employee_names))  # Remove duplicates

def extract_with_mapping(employee_name, session_dir, filename_mapping):
    """Extract shifts for specific employee using filename mapping"""
    results = []
    
    # First pass: Extract all employee names from all files
    all_employee_names = set()
    for secure_fname in os.listdir(session_dir):
        if not secure_fname.endswith('.docx'):
            continue
            
        filepath = os.path.join(session_dir, secure_fname)
        
        try:
            doc = Document(filepath)
            
            # Find the table and extract all possible names
            for table in doc.tables:
                for row in table.rows[1:]:  # Skip header
                    shift_type = row.cells[0].text.strip()
                    
                    # Skip "Assenti" entries
                    if 'assenti' in shift_type.lower():
                        continue
                    
                    for cell in row.cells[1:]:
                        cell_text = cell.text.strip()
                        if cell_text and not cell_text.isspace():
                            # Simple extraction for building the name list
                            base_names = [name.strip() for name in cell_text.replace('\n', ',').split(',') if name.strip()]
                            
                            for base_name in base_names:
                                # Remove non-name parentheses content
                                clean_base = re.sub(r'\([^)]*(?:turno|shift|ore|h|:|\d+)[^)]*\)', '', base_name, flags=re.IGNORECASE)
                                
                                # Extract main name
                                main_name = re.sub(r'\([^)]*\)', '', clean_base).strip()
                                if main_name and len(main_name) > 1 and not is_roman_numeral(main_name):
                                    all_employee_names.add(main_name)
                                
                                # Extract potential names from parentheses
                                parentheses_matches = re.findall(r'\(([^)]+)\)', clean_base)
                                for match in parentheses_matches:
                                    potential_name = match.strip()
                                    
                                    # Skip Roman numerals
                                    if is_roman_numeral(potential_name):
                                        continue
                                        
                                    if (potential_name and 
                                        len(potential_name) > 1 and 
                                        not re.match(r'^[\d:.-]+$', potential_name) and
                                        not re.search(r'\b(?:turno|shift|ore|h)\b', potential_name, re.IGNORECASE)):
                                        all_employee_names.add(potential_name)
        
        except Exception as e:
            print(f"Error in name extraction for {secure_fname}: {e}")
            continue
    
    # Second pass: Extract shifts for the specific employee
    for secure_fname in os.listdir(session_dir):
        if not secure_fname.endswith('.docx'):
            continue
            
        filepath = os.path.join(session_dir, secure_fname)
        original_filename = filename_mapping.get(secure_fname, secure_fname)
        
        try:
            doc = Document(filepath)
            
            # Extract date range from ORIGINAL filename
            start_day, start_month, end_day, end_month, start_year, end_year = extract_employee_shifts.extract_date_range_from_filename(original_filename)
            if not all([start_day, start_month, end_day, end_month]):
                continue
                
            # Get actual dates for this week
            week_dates = extract_employee_shifts.get_week_dates_from_range(start_day, start_month, end_day, end_month, start_year, end_year)
            
            # Find the table
            for table in doc.tables:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                days = headers[1:]
                
                for row in table.rows[1:]:
                    shift_type = row.cells[0].text.strip()
                    
                    for i, cell in enumerate(row.cells[1:]):
                        cell_text = cell.text.strip()
                        if cell_text and not cell_text.isspace():
                            # Extract all possible names using the enhanced method
                            employee_names = extract_employee_names_from_cell(cell_text, all_employee_names)
                            
                            # Check if our target employee is in this cell
                            if employee_name in employee_names:
                                if 'assenti' in shift_type.lower():
                                    continue
                                    
                                day_name = days[i] if i < len(days) else f'Day{i+1}'
                                date_str = week_dates[i] if i < len(week_dates) else ''
                                
                                results.append({
                                    'File': original_filename,
                                    'Data': date_str,
                                    'Giorno': day_name,
                                    'Turno': shift_type,
                                    'Dipendente': employee_name
                                })
                                
                                # Add weekend shifts for "Guardia" on Friday
                                if 'guardia' in shift_type.lower() and day_name.lower() == 'venerdì':
                                    # Add Saturday
                                    saturday_date = extract_employee_shifts.add_days_to_date(date_str, 1) if date_str else ''
                                    results.append({
                                        'File': original_filename,
                                        'Data': saturday_date,
                                        'Giorno': 'Sabato',
                                        'Turno': shift_type,
                                        'Dipendente': employee_name
                                    })
                                    
                                    # Add Sunday
                                    sunday_date = extract_employee_shifts.add_days_to_date(date_str, 2) if date_str else ''
                                    results.append({
                                        'File': original_filename,
                                        'Data': sunday_date,
                                        'Giorno': 'Domenica',
                                        'Turno': shift_type,
                                        'Dipendente': employee_name
                                    })
        
        except Exception as e:
            print(f"Error processing {original_filename}: {e}")
            continue
    
    return results

@app.route('/download/<session_id>/<path:filename>', methods=['GET', 'HEAD'])
def download_file(session_id, filename):
    """Primary download endpoint (supports browser and fetch)."""
    safe_name = secure_filename(filename)
    session_dir = os.path.join(UPLOAD_FOLDER, session_id)
    if not os.path.isdir(session_dir):
        return jsonify({'error': 'Session not found'}), 404
    file_path = os.path.join(session_dir, safe_name)
    if not os.path.isfile(file_path):
        return jsonify({'error': 'File not found'}), 404
    try:
        # Ensure file is ready (size > 0)
        for _ in range(50):
            if os.path.getsize(file_path) > 0:
                break
            time.sleep(0.02)
        file_size = os.path.getsize(file_path)
        resp = send_file(
            file_path,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=safe_name,
            conditional=False
        )
        # Strengthen headers
        resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        resp.headers['Content-Length'] = str(file_size)
        resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        resp.headers['Pragma'] = 'no-cache'
        resp.headers['Expires'] = '0'
        resp.headers['X-Content-Type-Options'] = 'nosniff'
        resp.headers['Content-Disposition'] = f'attachment; filename={safe_name}; filename*=UTF-8''{safe_name}'
        return resp
    except Exception as e:
        return jsonify({'error': f'Download failed: {str(e)}'}), 500

@app.route('/api/download/<session_id>/<path:filename>', methods=['GET'])
def api_download_file(session_id, filename):
    """Raw bytes endpoint for robust fetch()+blob download inside webview."""
    safe_name = secure_filename(filename)
    session_dir = os.path.join(UPLOAD_FOLDER, session_id)
    if not os.path.isdir(session_dir):
        return jsonify({'error': 'Session not found'}), 404
    file_path = os.path.join(session_dir, safe_name)
    if not os.path.isfile(file_path):
        return jsonify({'error': 'File not found'}), 404
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        resp = make_response(data)
        resp.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        resp.headers['Content-Length'] = str(len(data))
        resp.headers['Cache-Control'] = 'no-store'
        resp.headers['Content-Disposition'] = f'attachment; filename={safe_name}'
        resp.headers['X-Content-Type-Options'] = 'nosniff'
        return resp
    except Exception as e:
        return jsonify({'error': f'API download failed: {str(e)}'}), 500

@app.route('/cleanup/<session_id>', methods=['POST'])
def cleanup_session(session_id):
    try:
        session_dir = os.path.join(UPLOAD_FOLDER, session_id)
        if os.path.exists(session_dir):
            shutil.rmtree(session_dir)
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'error': f'Cleanup failed: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5001)

# Lightweight API (not used in UI yet) to dump a simple PNG of summary counts
@app.route('/api/heatmap_png', methods=['POST'])
def heatmap_png():
    try:
        data = request.get_json(force=True)
        summary = data.get('summary', {}) if isinstance(data, dict) else {}
        if not summary:
            return jsonify({'error': 'No summary provided'}), 400
        if Image is None:
            return jsonify({'error': 'Pillow not available'}), 500
        # Collect shift types
        shift_types = set()
        for shifts in summary.values():
            for s in shifts.keys():
                shift_types.add(s)
        shift_types = sorted(shift_types)
        cell_w = 120
        cell_h = 40
        header_h = 50
        left_w = 180
        rows = len(summary)
        cols = len(shift_types)+2  # employee + total
        img_w = left_w + cell_w*(cols-1)
        img_h = header_h + cell_h*rows
        img = Image.new('RGB', (img_w, img_h), 'white')
        draw = ImageDraw.Draw(img)
        try:
            font = ImageFont.load_default()
        except Exception:
            font = None
        # Headers
        draw.rectangle([0,0,img_w,header_h], fill=(102,126,234))
        draw.text((10, header_h/2), 'Employee', fill='white', anchor='lm', font=font)
        for idx, st in enumerate(shift_types):
            x = left_w + idx*cell_w + 10
            draw.text((x, header_h/2), st[:12], fill='white', anchor='lm', font=font)
        draw.text((left_w + len(shift_types)*cell_w + 10, header_h/2), 'Total', fill='white', anchor='lm', font=font)
        # Rows
        for r,(emp, shifts) in enumerate(sorted(summary.items())):
            y0 = header_h + r*cell_h
            draw.rectangle([0,y0,left_w,y0+cell_h], fill=(248,249,250))
            draw.text((10, y0+cell_h/2), emp[:20], fill='black', anchor='lm', font=font)
            total = 0
            for c, st in enumerate(shift_types):
                val = shifts.get(st,0)
                total += val
                intensity = min(val/10,1.0)
                base = (102,126,234)
                bg = tuple(int(b* (0.3+0.7*intensity)) for b in base) if val>0 else (248,249,250)
                x0 = left_w + c*cell_w
                draw.rectangle([x0,y0,x0+cell_w,y0+cell_h], fill=bg)
                if val:
                    draw.text((x0+cell_w/2, y0+cell_h/2), str(val), fill='white' if val>5 else 'black', anchor='mm', font=font)
            # total cell
            x0 = left_w + len(shift_types)*cell_w
            draw.rectangle([x0,y0,x0+cell_w,y0+cell_h], fill=(40,167,69))
            draw.text((x0+cell_w/2, y0+cell_h/2), str(total), fill='white', anchor='mm', font=font)
        # Grid lines (optional)
        for r in range(rows+1):
            y = header_h + r*cell_h
            draw.line([0,y,img_w,y], fill='#e0e0e0')
        for c in range(cols):
            x = (left_w if c>0 else 0) + (c-1)*cell_w if c>0 else 0
            if c==0:
                draw.line([left_w,0,left_w,img_h], fill='#e0e0e0')
            else:
                draw.line([left_w + (c-1)*cell_w,0,left_w + (c-1)*cell_w,img_h], fill='#e0e0e0')
        bio = BytesIO()
        img.save(bio, format='PNG')
        bio.seek(0)
        inline = request.args.get('inline') == '1'
        return send_file(bio, mimetype='image/png', as_attachment=not inline, download_name='heatmap.png')
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/heatmap_save', methods=['POST'])
def heatmap_save():
    """Save heatmap PNG directly to user's Downloads without navigating UI."""
    if Image is None:
        return jsonify({'error': 'Pillow not available'}), 500
    try:
        data = request.get_json(force=True)
        summary = data.get('summary', {}) if isinstance(data, dict) else {}
        # Build shift_types from summary (alphabetical ordering like client)
        shift_types = sorted({s for v in summary.values() for s in v.keys()})
        employee_names = sorted(summary.keys())
        max_emp_len = max((len(e) for e in employee_names), default=10)
        max_shift_len = max((len(s) for s in shift_types), default=5) if shift_types else 5
        # Approximate char width in pixels (increased for larger fonts)
        left_w = max(200, max_emp_len * 11 + 40)
        cell_w = max(110, min(200, max_shift_len * 11 + 40))
        cell_h = 65  # Increased height for larger fonts
        header_h = 70  # Increased header height
        rows = len(employee_names)
        cols = len(shift_types) + 2  # shifts + total
        img_w = left_w + cell_w * (cols - 1)
        img_h = header_h + cell_h * rows
        img = Image.new('RGB', (img_w, img_h), 'white')
        draw = ImageDraw.Draw(img)

        # Fonts
        font = None
        font_small = None
        value_font = None
        try:
            font_paths = [
                '/System/Library/Fonts/Supplemental/Arial.ttf',
                '/System/Library/Fonts/Supplemental/Helvetica.ttc',
                '/System/Library/Fonts/Supplemental/Verdana.ttf'
            ]
            for fp in font_paths:
                if os.path.exists(fp):
                    # Header / labels (increased font sizes)
                    font = ImageFont.truetype(fp, 22)  # Increased from 18
                    font_small = ImageFont.truetype(fp, 20)  # Increased from 16
                    # Much larger numeric values font for better readability
                    try:
                        value_font = ImageFont.truetype(fp, 32)  # Increased from 22
                    except Exception:
                        value_font = font
                    break
        except Exception:
            pass
        if font is None:
            try:
                font = ImageFont.load_default()
                font_small = font
                value_font = font
            except Exception:
                font = None
                font_small = None
                value_font = None

        # Determine max value for normalization (avoid division by zero)
        all_vals = [v for shifts in summary.values() for v in shifts.values() if isinstance(v, (int, float))]
        max_val = max(all_vals) if all_vals else 1

        base_col = (102,126,234)      # #667eea
        light_col = (238,242,255)     # very light lavender
        total_col = (40,167,69)       # green
        header_col = base_col
        alt_row = (248,249,250)

        # Helper function for thick outlined text
        def draw_outlined_text(x, y, text, font_obj, fill_color, outline_color=(0,0,0), outline_width=3):
            """Draw text with thick outline for better readability"""
            if font_obj is None:
                draw.text((x, y), text, fill=fill_color, anchor='mm')
                return
            
            # Draw outline by drawing text at multiple offset positions
            for dx in range(-outline_width, outline_width + 1):
                for dy in range(-outline_width, outline_width + 1):
                    if dx == 0 and dy == 0:
                        continue
                    draw.text((x + dx, y + dy), text, fill=outline_color, anchor='mm', font=font_obj)
            
            # Draw the main text on top
            draw.text((x, y), text, fill=fill_color, anchor='mm', font=font_obj)

        # Header background
        draw.rectangle([0, 0, img_w, header_h], fill=header_col)
        # Column headers
        draw.text((10, header_h/2), 'Employee', fill='white', anchor='lm', font=font)
        for idx, st in enumerate(shift_types):
            x = left_w + idx * cell_w + 10
            label = st
            if len(label) > 16:
                label = label[:15] + '…'
            draw.text((x, header_h/2), label, fill='white', anchor='lm', font=font_small or font)
        draw.text((left_w + (len(shift_types)) * cell_w + 10, header_h/2), 'Total', fill='white', anchor='lm', font=font)

        # Rows and cells
        for r, emp in enumerate(employee_names):
            y0 = header_h + r * cell_h
            # Employee name band
            draw.rectangle([0, y0, left_w, y0 + cell_h], fill=alt_row)
            name_label = emp
            if len(name_label) > 24:
                name_label = name_label[:23] + '…'
            draw.text((10, y0 + cell_h/2), name_label, fill='black', anchor='lm', font=font_small or font)
            shifts = summary[emp]
            total = 0
            for c, st in enumerate(shift_types):
                val = shifts.get(st, 0)
                total += val
                intensity = (val / max_val) if max_val else 0
                if val:
                    # Interpolate color
                    bg = tuple(int(light_col[i] + (base_col[i] - light_col[i]) * intensity) for i in range(3))
                else:
                    bg = alt_row
                x0 = left_w + c * cell_w
                draw.rectangle([x0, y0, x0 + cell_w, y0 + cell_h], fill=bg)
                if val:
                    # Decide text color based on background brightness for contrast
                    brightness = (bg[0]*299 + bg[1]*587 + bg[2]*114) / 1000
                    txt_col = 'white' if brightness < 160 else 'black'
                    outline_col = (0,0,0) if txt_col == 'white' else (255,255,255)
                    
                    # Use outlined text for much better readability
                    cx = x0 + cell_w/2
                    cy = y0 + cell_h/2
                    draw_outlined_text(cx, cy, str(val), value_font or font_small or font, txt_col, outline_col, outline_width=3)
            # Total cell
            x0 = left_w + len(shift_types) * cell_w
            draw.rectangle([x0, y0, x0 + cell_w, y0 + cell_h], fill=total_col)
            # Total value with thick outline for maximum readability
            cx_tot = x0 + cell_w/2
            cy_tot = y0 + cell_h/2
            draw_outlined_text(cx_tot, cy_tot, str(total), value_font or font, 'white', (0,0,0), outline_width=3)

        # Grid lines
        for r in range(rows + 1):
            y = header_h + r * cell_h
            draw.line([0, y, img_w, y], fill='#d0d5dd')
        # Vertical lines (after employee col)
        draw.line([left_w, 0, left_w, img_h], fill='#d0d5dd')
        for c in range(len(shift_types)):
            x = left_w + c * cell_w
            draw.line([x, 0, x, img_h], fill='#e2e6ef')
        draw.line([img_w-1, 0, img_w-1, img_h], fill='#d0d5dd')
        # Save to downloads
        try:
            downloads_dir = get_user_downloads_dir()
        except Exception:
            downloads_dir = os.path.expanduser('~')
        os.makedirs(downloads_dir, exist_ok=True)
        base_name = 'heatmap.png'
        target = os.path.join(downloads_dir, base_name)
        if os.path.exists(target):
            stem, ext = os.path.splitext(base_name)
            counter = 1
            while True:
                cand = os.path.join(downloads_dir, f"{stem} ({counter}){ext}")
                if not os.path.exists(cand):
                    target = cand
                    break
                counter += 1
        img.save(target, format='PNG')
        return jsonify({'success': True, 'saved_path': target})
    except Exception as e:
        return jsonify({'error': str(e)}), 500
